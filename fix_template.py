from docx import Document
import os

print("Starting fix_template.py script...")
template_path = "my_template.docx.docx"
if os.path.exists(template_path):
    print(f"Template {template_path} found.")
    doc = Document(template_path)
    # The user input said "-空白模板" but the inspection showed "- 空白模板"
    # Let's try to match both
    
    found_count = 0
    for paragraph in doc.paragraphs:
        if "单人候选人面试评估报告" in paragraph.text and "模板" in paragraph.text:
            print(f"Found potential target paragraph: {paragraph.text}")
            # Use regex or simple split/replace
            if "- 空白模板" in paragraph.text:
                paragraph.text = paragraph.text.replace("- 空白模板", "").strip()
                found_count += 1
            elif "-空白模板" in paragraph.text:
                paragraph.text = paragraph.text.replace("-空白模板", "").strip()
                found_count += 1
            elif "空白模板" in paragraph.text:
                paragraph.text = paragraph.text.replace("空白模板", "").strip()
                # Also remove trailing dashes
                paragraph.text = paragraph.text.rstrip("- ").strip()
                found_count += 1
    
    if found_count > 0:
        doc.save(template_path)
        print(f"Successfully updated template. Removed template suffix from {found_count} paragraphs.")
    else:
        print("Could not find the target text with exact match. Trying more flexible match...")
        for paragraph in doc.paragraphs:
            if "单人候选人面试评估报告" in paragraph.text:
                print(f"Paragraph contains '单人候选人面试评估报告': {paragraph.text}")
                if "模板" in paragraph.text:
                    # Replace anything after the closing parenthesis or similar
                    # For now, just replace the exact substrings if possible
                    pass
else:
    print(f"Template file {template_path} not found.")
