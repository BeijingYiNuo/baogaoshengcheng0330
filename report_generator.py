# report_generator.py
import os
import json
from openai import AsyncOpenAI
from prompts import PROMPT_SYSTEM, ATTENTION_PROMPT, PROMPT_IMGCODE
from doc2md import split_md
from dotenv import load_dotenv
from docx import Document as DocumentParse
from docx.text.paragraph import Paragraph
import re
import asyncio
from state_mgr import StateManager, FileInfo
from io import BytesIO
import textwrap
from docx.shared import Inches
from datetime import datetime
import random

load_dotenv()


def record(prompt: str, response: str, skip: bool = True):
    if skip:
        return
    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
    ts = f"{ts}_{random.randint(100,999)}"
    os.makedirs("record", exist_ok=True)
    open(f"record/{ts}.md", "w", encoding="utf-8").write(
        f"# PROMPT\n{prompt}\n# RESPONSE\n{response}"
    )


def strip_code_fence(text):
    return re.sub(r"```.*?\n|```", "", text, flags=re.DOTALL)


def remove_invalid_key(doc: str, data: dict):
    keys = set[str](data.keys())
    for key in keys:
        if "{{" + key in doc:
            continue
        else:
            del data[key]


def build_and_run_draw(
    code_str: str,
    draw_code: str = open(
        os.path.dirname(__file__) + "/draw.py", encoding="utf-8"
    ).read(),
) -> BytesIO:
    code_str = strip_code_fence(code_str)
    code_str = textwrap.indent(code_str, " " * 8)
    pattern = "#<REPLACE_START>.*?#<REPLACE_END>"
    replacement = "#<REPLACE_START>\n" + code_str + "\n    #<REPLACE_END>"
    new_script = re.sub(pattern, replacement, draw_code, flags=re.S)
    local_env = {}
    exec(new_script, {}, local_env)
    draw_func = local_env["draw"]
    img = draw_func()
    return img


def get_grade(score):
    if score >= 90:
        return "优秀"
    elif score >= 80:
        return "良好"
    elif score >= 70:
        return "合格"
    else:
        return "不合格"


def score_postprocess(all_data: dict[str, str | float]):
    score_dim = dict[str, float]()
    level_dim = dict[str, str]()
    for key, value in all_data.items():
        if key.startswith("Score_AI_"):
            try:
                value = float(value)
            except:
                value = 70.0
            dim_name = key.replace("Score_AI_", "")
            score_dim[dim_name] = value
            level_dim[dim_name] = get_grade(value)
    for dim_name in level_dim:
        all_data["Score_Level_" + dim_name] = level_dim[dim_name]

    if score_dim:
        total_avg = round(sum(score_dim.values()) / len(score_dim), 1)
    else:
        total_avg = 0.0
    all_data["Score_TotalAvg"] = total_avg
    all_data["Score_Comparison"] = "高于" if total_avg > 70 else "低于"
    all_data["Score_DimCount"] = len(level_dim)
    all_data["Score_MetCount"] = sum(
        1 for _, dim_score in score_dim.items() if dim_score > 70
    )
    all_data["Score_UnmetCount"] = (
        all_data["Score_DimCount"] - all_data["Score_MetCount"]
    )


async def draw_chart(
    client: AsyncOpenAI,
    model: str,
    shadow_doc: str,
    data: dict[str, BytesIO],
    draw_code: str = open(
        os.path.dirname(__file__) + "/draw.py", encoding="utf-8"
    ).read(),
) -> None:
    prompt = PROMPT_IMGCODE.replace("{CONTEXT}", shadow_doc).replace(
        "{DRAW_CODE}", draw_code
    )

    matches = re.findall(r"\[\[(Chart_\w+)", prompt)

    async def process_match(match: str):
        m_prompt = prompt.replace("{KEY}", match)

        response = await client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": m_prompt}],
            temperature=0.3,
        )
        record(m_prompt, response.choices[0].message.content)
        result = build_and_run_draw(code_str=response.choices[0].message.content)
        return match, result

    # 并发执行
    tasks = [process_match(m) for m in matches]
    results = await asyncio.gather(*tasks)

    # 汇总写回 data
    for match, result in results:
        data[match] = result


async def generate_interview_data_multi_stage_async(
    api_key: str,
    base_url: str,
    model: str,
    template_md: str,
    resume_text: str,
    transcript_text: str,
    all_data: dict[str, str | float | int] = {},
    sm: StateManager = None,
) -> FileInfo:
    if sm is not None:
        f_info = sm.start()
    else:
        f_info = FileInfo.from_info(idx=0)

    client = AsyncOpenAI(api_key=api_key, base_url=base_url)
    templates = split_md(template_md)

    async def fetch_data_from_llm(idx, template_content):
        prompt = (
            PROMPT_SYSTEM.replace("{TEMPLATE}", template_content)
            .replace("{RESUME}", resume_text)
            .replace("{VOICE2TEXT}", transcript_text)
            .replace("{ATTENTION}", ATTENTION_PROMPT)
        )
        response = await client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        record(prompt, response.choices[0].message.content)
        data = json.loads(response.choices[0].message.content)
        remove_invalid_key(template_content, data)
        shadow_doc: str = template_content
        for key, value in data.items():
            shadow_doc = shadow_doc.replace("{{" + key + "}}", str(value))
        chart_pattern = r"\[\[Chart_\w+"
        if re.search(chart_pattern, shadow_doc):
            await draw_chart(client, model, shadow_doc, data)
        return data

    tasks = [
        asyncio.create_task(fetch_data_from_llm(idx, content))
        for idx, content in enumerate(templates)
    ]
    done_counter = 0
    total_counter = len(tasks) + 2 if sm is not None else 1
    for future in asyncio.as_completed(tasks):
        try:
            data = await future
            for k in data:
                if k not in all_data:
                    v = data[k]
                    if isinstance(v, str) and v.endswith("。"):
                        v = v[:-1]
                    all_data[k] = v
        finally:
            done_counter += 1
            if sm is not None:
                sm.update(f_info.idx, done_counter / total_counter)
    score_postprocess(all_data)
    done_counter += 1
    if sm is not None:
        sm.update(f_info.idx, done_counter / total_counter, "完成评分")
        sm.done(f_info.idx)
    return f_info


def organize_run(paras, start_str: str = "{{", end_str: str = "}}"):
    for para in paras:
        runs = para.runs
        if not runs:
            continue

        i = 0
        while i < len(runs):
            if start_str in runs[i].text:
                start_run = i
                merged_text = runs[i].text
                j = i
                while end_str not in merged_text and j + 1 < len(runs):
                    j += 1
                    merged_text += runs[j].text
                if end_str in merged_text and j != i:
                    end_pos = merged_text.rindex(end_str) + 2
                    full_placeholder = merged_text[:end_pos]
                    runs[start_run].text = full_placeholder
                    for k in range(start_run + 1, j + 1):
                        runs[k].text = merged_text[end_pos:] if k == j else ""
                    i = j
            i += 1


def replace_in_para(
    para: Paragraph,
    data: dict,
    pattern=re.compile(r"\{\{(\w+)(:[^}]*)?\}\}|\[\[(\w+)(:[^}]*)?\]\]"),
):
    for run in para.runs:
        text = run.text
        matches = list(pattern.finditer(text))
        if not matches:
            continue
        for m in matches:
            key = m.group(1) or m.group(3)
            value = data.get(key)
            if value is None:
                continue
            elif key.startswith("Chart_") and isinstance(value, BytesIO):
                run.text = run.text.replace(m.group(0), "")
                run.add_picture(value, Inches(6))
                break
            else:
                run.text = run.text.replace(m.group(0), str(value))


def fill_docx_template(template_path, output_path, data: dict):
    doc = DocumentParse(template_path)
    organize_run(doc.paragraphs)
    organize_run(doc.paragraphs, "[[", "]]")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                organize_run(cell.paragraphs)
                organize_run(cell.paragraphs, "[[", "]]")

    for para in doc.paragraphs:
        replace_in_para(para, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, data)
    doc.save(output_path)
    return output_path
