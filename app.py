# app.py
import streamlit as st
import os
from report_generator import generate_interview_data_multi_stage, fill_docx_template
from voice2text import transcribe_audio_generator
from docx import Document
from io import BytesIO
import pypdf
import json
import datetime
import tempfile
import subprocess

CONFIG_FILE = "config.json"
DEFAULT_TEMPLATE = "template.docx"
REPORTS_DIR = "reports"
TEMP_DIR = "temp_audio"

if not os.path.exists(REPORTS_DIR):
    os.makedirs(REPORTS_DIR, exist_ok=True)

if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR, exist_ok=True)

# 检查默认模板是否存在
if not os.path.exists(DEFAULT_TEMPLATE):
    # 自动创建一个简单的基础模板
    doc = Document()
    doc.add_heading("面试评估报告", 0)
    doc.add_paragraph("面试日期: {{InterviewDate}}")
    doc.add_paragraph("面试形式: {{InterviewFormat}}")
    doc.add_paragraph("面试评估官: {{Interviewer}}")
    doc.add_paragraph("报告编制人: {{ReportCompiler}}")
    doc.add_paragraph("姓名: {{CandidateName}}")
    doc.add_paragraph("职位: {{TargetPosition}}")
    doc.add_paragraph("结论: {{FinalDecision}}")
    doc.save(DEFAULT_TEMPLATE)


def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_config(config_dict):
    with open(CONFIG_FILE, "w") as f:
        json.dump(config_dict, f)


# 初始化配置
if "config" not in st.session_state:
    st.session_state.config = load_config()

if "last_report" not in st.session_state:
    st.session_state.last_report = None

if "transcription_text" not in st.session_state:
    st.session_state.transcription_text = ""

if "last_audio_file" not in st.session_state:
    st.session_state.last_audio_file = None

if "resume_text_cached" not in st.session_state:
    st.session_state.resume_text_cached = ""

if "last_resume_file" not in st.session_state:
    st.session_state.last_resume_file = None

if "transcript_text_cached" not in st.session_state:
    st.session_state.transcript_text_cached = ""

if "last_transcript_file" not in st.session_state:
    st.session_state.last_transcript_file = None

if "interview_result_summary" not in st.session_state:
    st.session_state.interview_result_summary = None


def extract_text_from_pdf(file):
    pdf_reader = pypdf.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def generate_summary_docx(summary_data):
    """
    生成一个只包含面试结果总览表的 Word 文档。
    """
    doc = Document()
    doc.add_heading("面试结果总览表", level=1)

    # 创建表格 2列 N行
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("候选人姓名", summary_data.get("CandidateName", "未知")),
        ("应聘岗位", summary_data.get("TargetPosition", "未知")),
        ("综合平均分", str(summary_data.get("TotalAvgScore", "0"))),
        ("面试通过情况", summary_data.get("PreScreeningResult", "未知")),
        ("备注", summary_data.get("OtherDetails", "无")),
    ]

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# 设置页面配置
st.set_page_config(page_title="面试报告生成器", layout="wide")

# 注入 CSS 隐藏右上角默认加载图标 (那个奥运图标)
st.markdown(
    """
    <style>
    /* 针对不同版本的 Streamlit 隐藏右上角菜单和页脚，但保留 Header 容器以便侧边栏按钮正常工作 */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* 仅仅隐藏 header 内部除侧边栏按钮以外的其他元素（如果需要更彻底的隐藏，可注释掉） */
    header[data-testid="stHeader"] {
        background-color: transparent !important;
        border: none !important;
    }
    
    /* 隐藏右上角 StatusWidget (加载图标) */
    div[data-testid="stStatusWidget"] {
        display: none !important;
    }
    
    /* 替代：在页面中心下方增加一个精致的脉冲加载条（仅在需要时通过 st.markdown 展示） */
    @keyframes pulse {
        0% { opacity: 0.5; transform: scaleX(0.95); }
        50% { opacity: 1; transform: scaleX(1); }
        100% { opacity: 0.5; transform: scaleX(0.95); }
    }
    .custom-loader {
        height: 4px;
        width: 100%;
        background-color: #007aff;
        border-radius: 2px;
        animation: pulse 1.5s infinite ease-in-out;
        margin: 10px 0;
    }
    </style>
""",
    unsafe_allow_html=True,
)

# 侧边栏配置 API
with st.sidebar:
    st.title("⚙️ 配置")

    # 获取默认值
    saved_api_key = st.session_state.config.get("api_key", "")
    saved_base_url = st.session_state.config.get(
        "base_url", "https://api.openai.com/v1"
    )
    saved_model = st.session_state.config.get("model", "gpt-4o")

    # 输入组件
    api_key = st.text_input(
        "API Key", value=saved_api_key, type="password", placeholder="输入您的 API Key"
    )
    base_url = st.text_input("Base URL", value=saved_base_url)

    model_list = [
        "gpt-4o",
        "gpt-4-turbo",
        "gpt-3.5-turbo",
        "deepseek-chat",
        "claude-3-5-sonnet-20240620",
    ]
    try:
        model_index = model_list.index(saved_model)
    except ValueError:
        model_index = 0
    model_name = st.selectbox("选择模型", model_list, index=model_index)

    # 保存配置按钮
    if st.button("💾 保存配置 (持久化)"):
        new_config = {"api_key": api_key, "base_url": base_url, "model": model_name}
        save_config(new_config)
        st.session_state.config = new_config
        st.success("✅ 配置已保存到本地文件，下次刷新页面不会丢失！")

    st.divider()
    st.subheader("🛠️ 系统状态")
    # API 状态
    if st.session_state.config.get("api_key"):
        st.success("✅ API 已配置")
    else:
        st.warning("⚠️ API 未配置")

    # FFmpeg 状态
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True)
        st.success("✅ FFmpeg 已就绪 (支持录音解析)")
    except:
        st.error("❌ FFmpeg 未找到 (录音功能受限)")

    st.divider()
    st.subheader("📚 历史文档")
    history_files = sorted(
        [f for f in os.listdir(REPORTS_DIR) if f.endswith(".docx")], reverse=True
    )
    if history_files:
        for f in history_files:
            # 获取文件修改时间并格式化
            file_path = os.path.join(REPORTS_DIR, f)
            mtime = os.path.getmtime(file_path)
            dt_str = datetime.datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M")

            # UI 展示：文件名 + 时间
            with st.expander(f"📄 {f}", expanded=False):
                st.caption(f"🕒 生成时间: {dt_str}")
                with open(file_path, "rb") as file:
                    st.download_button(
                        label="下载此报告",
                        data=file.read(),
                        file_name=f,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"sidebar_dl_{f}",
                        use_container_width=True,
                    )
    else:
        st.write("暂无历史文档")

    st.divider()
    st.subheader("📁 报告模板")
    current_template = DEFAULT_TEMPLATE
    st.info(f"当前使用默认模板: {DEFAULT_TEMPLATE}")


def render_step_one():
    st.title("📄 智能面试评估系统")
    st.markdown(
        """
        欢迎使用 **智能面试评估系统**。本工具通过 AI 技术深度分析面试录音及简历，自动生成专业且标准化的评估报告。
        请填写本次面试的基础信息，然后点击“下一步”开始生成报告。
    """
    )
    st.divider()

    with st.form("interview_info_form"):
        # 分为左右两列，使布局更饱满且更具“应用感”
        col_left, col_right = st.columns(2, gap="large")

        with col_left:
            st.markdown("#### 📅 面试基础信息")

            # 日期选择
            interview_date = st.date_input("面试日期", datetime.date.today())

            # 24小时制时间选择：左侧小时，右侧分钟
            st.write("**具体时间 (24h)**")
            time_col1, time_col2 = st.columns(2)
            with time_col1:
                h = st.selectbox("小时", [f"{i:02d}" for i in range(24)], index=10)
            with time_col2:
                m = st.selectbox(
                    "分钟", [f"{i:02d}" for i in range(0, 60, 10)], index=0
                )

            # 合并日期和时间
            full_datetime_str = f"{interview_date.strftime('%Y年%m月%d日')} {h}:{m}"

            # 面试形式 - 水平排列
            interview_format = st.radio(
                "📍 **面试形式**", ("线上", "线下"), index=0, horizontal=True
            )

        with col_right:
            st.markdown("#### 👥 人员配置")

            # 评估官
            interviewer = st.text_input(
                "👨‍🏫 **面试评估官**", placeholder="请输入评估官姓名", help="例如：张三"
            )

            # 编制人
            report_compiler = st.text_input(
                "📝 **报告编制人**",
                placeholder="请输入报告编制人姓名",
                help="例如：李四",
            )

            st.markdown("<br>", unsafe_allow_html=True)
            st.info("💡 **温馨提示**：\n\n以上信息将自动关联至报告模板的首页。")

        st.markdown("<br>", unsafe_allow_html=True)
        # 居中下一步按钮
        _, btn_col, _ = st.columns([1, 1, 1])
        with btn_col:
            submitted = st.form_submit_button(
                "🚀 下一步：上传文档", use_container_width=True
            )

        if submitted:
            if not all([interview_date, interviewer, report_compiler]):
                st.warning("⚠️ 请填写所有必填项！")
            else:
                # 保存表单数据到 session_state
                st.session_state.form_data = {
                    "InterviewTime": full_datetime_str,
                    "InterviewFormat": interview_format,
                    "InterviewerInfo": interviewer,
                    "ReportCreator": report_compiler,
                    "ReportCreateDate": datetime.date.today().strftime("%Y年%m月%d日"),
                }
                # 进入第二步
                st.session_state.step = 2
                st.rerun()


def render_step_two():
    # 主界面
    st.title("📝 面试报告生成")
    st.markdown(
        "输入面试全过程语音转换的文字和面试者简历，按照模板自动生成一份完整的面试报告。"
    )

    col_left, col_right = st.columns(2)

    with col_left:
        st.subheader("📄 面试者简历")

        # 文件上传逻辑
        uploaded_resume = st.file_uploader(
            "上传简历 (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"]
        )

        # 尝试从上传的文件中提取文本
        extracted_text = st.session_state.resume_text_cached
        if uploaded_resume:
            # 如果是新上传的文件，重新解析
            if st.session_state.last_resume_file != uploaded_resume.name:
                try:
                    if uploaded_resume.type == "application/pdf":
                        st.session_state.resume_text_cached = extract_text_from_pdf(
                            uploaded_resume
                        )
                    elif (
                        uploaded_resume.type
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        st.session_state.resume_text_cached = extract_text_from_docx(
                            uploaded_resume
                        )
                    else:  # txt
                        st.session_state.resume_text_cached = (
                            uploaded_resume.read().decode("utf-8")
                        )
                    st.session_state.last_resume_file = uploaded_resume.name
                    st.success("✅ 文件解析成功！")
                except Exception as e:
                    st.error(f"❌ 文件解析失败: {str(e)}")

            extracted_text = st.session_state.resume_text_cached

        # 文本框，如果解析成功则填充，否则允许手动输入
        resume_text = st.text_area(
            "简历内容",
            value=extracted_text if extracted_text else "",
            height=300,
            placeholder="解析后的内容将显示在这里，或者手动粘贴简历内容...",
        )

    with col_right:
        st.subheader("🎤 面试记录")

        # 增加一键填充示例数据的按钮
        if st.button("💡 填充演示示例数据", use_container_width=True):
            st.session_state.resume_text_cached = "姓名：张小明\n求职意向：高级前端开发工程师\n技能：React, Vue, Node.js, Webpack\n工作经历：某知名互联网大厂 5年经验，主导过千万级用户量项目的架构升级。"
            st.session_state.transcription_text = "面试官：请介绍一下你在 React 方面的项目经验。\n候选人：我曾经在项目中负责过整个前端组件库的搭建，通过引入 React.memo 和自定义 Hooks，将页面首屏加载速度提升了 40%。此外，我还解决了多个复杂的跨域和性能瓶颈问题..."
            st.rerun()

        # 面试记录文件上传
        uploaded_transcript = st.file_uploader(
            "上传面试记录 (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"]
        )

        # 语音文件上传 (新增)
        uploaded_audio = st.file_uploader(
            "上传面试录音 (MP3/WAV/M4A)", type=["mp3", "wav", "m4a", "flac"]
        )

        # 尝试从上传的文件中提取文本
        extracted_transcript = (
            st.session_state.transcription_text
            if st.session_state.transcription_text
            else st.session_state.transcript_text_cached
        )
        if uploaded_transcript:
            # 如果是新上传的文件，重新解析
            if st.session_state.last_transcript_file != uploaded_transcript.name:
                try:
                    if uploaded_transcript.type == "application/pdf":
                        st.session_state.transcript_text_cached = extract_text_from_pdf(
                            uploaded_transcript
                        )
                    elif (
                        uploaded_transcript.type
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        st.session_state.transcript_text_cached = (
                            extract_text_from_docx(uploaded_transcript)
                        )
                    else:  # txt
                        st.session_state.transcript_text_cached = (
                            uploaded_transcript.read().decode("utf-8")
                        )
                    st.session_state.last_transcript_file = uploaded_transcript.name
                    st.success("✅ 面试记录解析成功！")
                except Exception as e:
                    st.error(f"❌ 面试记录解析失败: {str(e)}")

            extracted_transcript = st.session_state.transcript_text_cached

        # 语音转文字逻辑 (新增)
        if uploaded_audio:
            # 如果是新上传的文件，清除之前的识别内容
            if st.session_state.last_audio_file != uploaded_audio.name:
                st.session_state.transcription_text = ""
                st.session_state.last_audio_file = uploaded_audio.name

            # 仅当没有识别过，或者用户点击“重新识别”按钮时才运行
            col_btn1, col_btn2 = st.columns([1, 1])
            with col_btn1:
                start_transcribe = st.button(
                    "🎙️ 开始语音转文字", type="secondary", use_container_width=True
                )
            with col_btn2:
                re_transcribe = st.button(
                    "🔄 重新识别", type="secondary", use_container_width=True
                )

            if start_transcribe or (re_transcribe and uploaded_audio):
                with st.status(
                    "🎧 正在解析语音并转换为文字...", expanded=True
                ) as status:
                    try:
                        # 将上传的文件保存为临时文件
                        with tempfile.NamedTemporaryFile(
                            delete=False,
                            suffix=f".{uploaded_audio.name.split('.')[-1]}",
                            dir=TEMP_DIR,
                        ) as tmp_file:
                            # 使用 getvalue() 获取全部内容，避免 read() 指标偏移问题
                            tmp_file.write(uploaded_audio.getvalue())
                            temp_audio_path = tmp_file.name

                        # 创建进度条
                        progress_bar = st.progress(0, text="准备识别...")

                        # 调用转文字生成器
                        all_text_container = st.empty()  # 用于动态展示识别出的文字

                        for progress, partial_text in transcribe_audio_generator(
                            temp_audio_path
                        ):
                            progress_bar.progress(
                                progress, text=f"识别中: {int(progress*100)}%"
                            )
                            # 实时展示完整识别出的文字
                            all_text_container.info(
                                f"🗨️ 已识别文字内容：\n\n{partial_text}"
                            )
                            st.session_state.transcription_text = partial_text

                        # 删除临时文件
                        os.unlink(temp_audio_path)
                        status.update(
                            label="✅ 语音转文字成功！",
                            state="complete",
                            expanded=False,
                        )
                    except Exception as e:
                        status.update(
                            label=f"❌ 语音转文字失败: {str(e)}", state="error"
                        )
                        st.error(f"❌ 语音转文字失败: {str(e)}")

            # 如果已经有识别结果，将其填充到提取的文本中
            if st.session_state.transcription_text:
                extracted_transcript = st.session_state.transcription_text

        transcript_text = st.text_area(
            "面试记录内容",
            value=extracted_transcript if extracted_transcript else "",
            height=300,
            placeholder="解析后的面试记录将显示在这里，或者手动粘贴面试过程记录...",
        )

    # 生成按钮
    if st.button("🚀 开始生成报告", type="primary"):
        # 清除上一次的报告状态，避免干扰本次生成
        st.session_state.last_report = None

        if not st.session_state.config.get("api_key"):
            st.error("❌ 请在侧边栏配置 API Key")
        elif not resume_text or not transcript_text:
            st.warning("⚠️ 请输入简历和面试记录内容")
        else:
            # 使用自定义的加载动画替代默认奥运图标
            loading_container = st.empty()

            try:
                # 阶段 1
                with loading_container:
                    st.markdown(
                        """
                        <div style="text-align: center; margin: 20px 0;">
                            <p style="color: #007aff; font-weight: 500; font-size: 1.1em;">🔍 阶段 1/3: 正在分析候选人简历并提取核心背景...</p>
                            <div class="custom-loader"></div>
                        </div>
                    """,
                        unsafe_allow_html=True,
                    )

                # 这里我们稍微调整一下逻辑，模拟分步效果，或者通过 LLM 的不同函数调用
                # 实际上 generate_interview_data_multi_stage 内部已经是分阶段的，
                # 但为了演示效果，我们可以在这里通过 st.status 或分步调用展示

                # 阶段 2
                with loading_container:
                    st.markdown(
                        """
                        <div style="text-align: center; margin: 20px 0;">
                            <p style="color: #007aff; font-weight: 500; font-size: 1.1em;">🎤 阶段 2/3: 正在深度解析面试录音，匹配岗位需求...</p>
                            <div class="custom-loader"></div>
                        </div>
                    """,
                        unsafe_allow_html=True,
                    )

                interview_data = generate_interview_data_multi_stage(
                    api_key=st.session_state.config.get("api_key"),
                    base_url=st.session_state.config.get("base_url"),
                    model=st.session_state.config.get("model"),
                    resume_text=resume_text,
                    transcript_text=transcript_text,
                )

                # 阶段 3
                with loading_container:
                    st.markdown(
                        """
                        <div style="text-align: center; margin: 20px 0;">
                            <p style="color: #007aff; font-weight: 500; font-size: 1.1em;">✍️ 阶段 3/3: 正在整合评估建议，填充 Word 报告模板...</p>
                            <div class="custom-loader"></div>
                        </div>
                    """,
                        unsafe_allow_html=True,
                    )

                # 将第一步的表单数据合并进来
                interview_data.update(st.session_state.form_data)

                # 2. 填充到 Word 模板
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                candidate_name = interview_data.get("CandidateName", "未知候选人")
                output_filename = f"面试报告_{candidate_name}_{timestamp}.docx"
                output_docx_path = os.path.join(REPORTS_DIR, output_filename)

                fill_docx_template(current_template, output_docx_path, interview_data)

                # 保存到 session_state
                st.session_state.last_report = {
                    "filename": output_filename,
                    "path": output_docx_path,
                }

                # 提取并保存面试结果总览表数据
                st.session_state.interview_result_summary = {
                    "CandidateName": interview_data.get("CandidateName", "未知候选人"),
                    "TargetPosition": interview_data.get("TargetPosition", "未知职位"),
                    "TotalAvgScore": interview_data.get("TotalAvgScore", "0"),
                    "PreScreeningResult": interview_data.get(
                        "PreScreeningResult", "未知结论"
                    ),
                    "OtherDetails": interview_data.get("OtherDetails", "无"),
                }

                # 移除加载动画
                loading_container.empty()
                # 强制刷新侧边栏历史列表
                st.rerun()

            except Exception as e:
                loading_container.empty()
                st.error(f"❌ 生成报告失败: {str(e)}")

    # 显示最新生成的报告下载按钮和面试结果总览（如果存在）
    if st.session_state.last_report:
        st.divider()
        st.success(
            f"✅ 深度报告生成成功！已保存至历史文档：{st.session_state.last_report['filename']}"
        )

        # 结果总览展示
        if st.session_state.interview_result_summary:
            st.subheader("📊 面试结果总览")
            summary = st.session_state.interview_result_summary

            # 使用列表/表格形式展示
            col_sum1, col_sum2 = st.columns([1, 1])
            with col_sum1:
                st.write(f"**候选人姓名:** {summary['CandidateName']}")
                st.write(f"**应聘岗位:** {summary['TargetPosition']}")
                st.write(f"**综合平均分:** {summary['TotalAvgScore']}")
            with col_sum2:
                st.write(f"**面试通过情况:** {summary['PreScreeningResult']}")
                st.write(f"**备注:** {summary['OtherDetails']}")

        st.subheader("💾 导出选项")
        col_down1, col_down2 = st.columns(2)
        with col_down1:
            with open(st.session_state.last_report["path"], "rb") as f:
                st.download_button(
                    label="⬇️ 下载完整面试报告 (Word)",
                    data=f.read(),
                    file_name=st.session_state.last_report["filename"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="main_download_btn",
                    use_container_width=True,
                )
        with col_down2:
            if st.session_state.interview_result_summary:
                summary_docx = generate_summary_docx(
                    st.session_state.interview_result_summary
                )
                st.download_button(
                    label="⬇️ 单独导出面试结果总览表",
                    data=summary_docx,
                    file_name=f"面试结果总览_{st.session_state.interview_result_summary['CandidateName']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="summary_download_btn",
                    use_container_width=True,
                )


# 页面路由
if "step" not in st.session_state:
    st.session_state.step = 1

if st.session_state.step == 1:
    render_step_one()
else:
    render_step_two()

# 页脚
st.divider()
st.caption("© 2026 面试评估系统 - 基于 LLM 构建")
