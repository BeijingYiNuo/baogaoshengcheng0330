from docx import Document
import os

template_path = "my_template.docx.docx"
if os.path.exists(template_path):
    doc = Document(template_path)
    print(f"Reading {template_path}:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    
    for i, t in enumerate(doc.tables):
        print(f"Table {i}:")
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"  T{i}R{r_idx}C{c_idx}: {cell.text.strip()}")
else:
    print(f"File {template_path} not found.")
