
import os
import shutil
from docx import Document

# 定义替换规则，将【】中的描述性文本映射到系统使用的 {{Key}}
replacements = {
    "【填写岗位】": "{{TargetPosition}}",
    "【填写应聘岗位，例：电话销售岗】": "{{AppliedPosition}}",
    "【填写应聘岗位】": "{{AppliedPosition}}",
    "【填写候选人姓名】": "{{CandidateName}}",
    "【填写面试时间，例：2026年03月XX日 XX:XX-XX:XX】": "{{InterviewTime}}",
    "【填写面试形式，例：线上面试/线下面试】": "{{InterviewFormat}}",
    "【填写评估官信息，例：HR：XXX；销售负责人：XXX】": "{{InterviewerInfo}}",
    "【填写编制人姓名】": "{{ReportCreator}}",
    "【填写编制时间，例：2026年03月XX日】": "{{ReportCreateDate}}",
    # 1.1 候选人基础档案 - 表格内替换
    "年龄": "年龄：{{Age}}",
    "性别": "性别：{{Gender}}",
    "学历/专业": "学历/专业：{{Education}}",
    "工作年限/电销经验": "工作年限/电销经验：{{YearsOfExperience}} / {{HasSalesExp}}",
    "过往背景": "过往背景：{{Background}}",
    "求职薪资": "求职薪资：{{ExpectedSalary}}",
    "意向岗位": "意向岗位：{{IntendedPosition}}",
    "【填写核心亮点，例：具备X年 电销经验，月均成单X单，熟悉电销话术】": "{{ResumeHighlights}}",
    "【已核实/未提及/存疑】": "{{VerificationStatus}}",
    "【填写核实内容】": "{{VerificationDetails}}",
    "【通过/未通过】": "{{PreScreeningResult}}",
    "【填写初筛依据】": "{{PreScreeningReason}}",
    "【填写综合平均分】": "{{TotalAvgScore}}",
    "【高于/低于/等于】": "{{ScoreComparison}}",
    "【完全达到/基本达到/未达到】": "{{GoalAchievement}}",
    "达标【填写数量】个": "达标{{MetCount}}个",
    "不合格【填写数量】个": "不合格{{UnmetCount}}个",
    "【填写达标情况】": "{{BasicCommStatus}}",
    "【填写优势维度】": "{{StrengthDimensions}}",
    "【填写薄弱维度】": "{{WeakDimensions}}",
    "原因【填写原因】": "原因{{Reason}}",
    "【高度一致/基本一致/偏差较大】": "{{ScoreConsistency}}",
    "偏差原因【填写原因】": "偏差原因{{ConsistencyReason}}",
    "【填写候选人表达状态，例：普通话标准，表达流畅，逻辑清晰/有轻微方言， 偶有卡顿】": "{{BasicCommStatusDesc}}",
    "匹配结论：【完全匹配/基本匹配/未匹配】": "匹配结论：{{BasicCommMatch}}",
    "是否满足电销沟通基础：【是/否】": "是否满足电销沟通基础：{{SatisfiesBasicComm}}",
    "待提升 点：【填写待提升点/无】": "待提升点：{{BasicCommImprovement}}",
    "【填写表现，例：有X年电销经验，熟悉话术，有明确业绩目标/无电销经验，学 习意愿强】": "{{SalesAbilityDesc}}",
    "【填写表现，例：面对客户拒绝能从容应对，接受电销岗位压力/面对追问略显 紧张，抗压性待提升】": "{{StressHandlingDesc}}",
    "是否具备电销抗压心态：【是/否】": "是否具备电销抗压心态：{{HasSalesMindset}}",
    "【填写表现，例：善于倾听，表达有亲和力，具备一定说服性/应答有偏差，亲 和力不足】": "{{CommExpressionDesc}}",
    "亮点/短板：【填写亮点/短板/ 无】": "亮点/短板：{{CommExpressionDetail}}",
    "【填写表现，例：求职意向坚定，职业规划贴合销售岗，无频繁离职/离职原因 不合理，意向不明确】": "{{StabilityDesc}}",
    "职业稳定性【高/中/低】": "职业稳定性{{StabilityLevel}}",
    "是否适合长期发展：【是/否】": "是否适合长期发展：{{IsLongTermFit}}",
    "核心考量：【填写 考量因素】": "核心考量：{{StabilityReason}}",
    "【基础沟通类】：【填写优势/无】": "【基础沟通类】：{{Strength_Comm}}",
    "【专业销售类】：【填写优势/无】": "【专业销售类】：{{Strength_Sales}}",
    "【抗压心态类】：【填写优势/无】": "【抗压心态类】：{{Strength_Stress}}",
    "【经验背景类】：【填写优势/无】": "【经验背景类】：{{Strength_Back}}",
    "【其他类】：【填写优势/无】": "【其他类】：{{Strength_Other}}",
    "【基础沟通类】：【填写劣势/无】": "【基础沟通类】：{{Weakness_Comm}}",
    "【专业销售类】：【填写劣势/无】": "【专业销售类】：{{Weakness_Sales}}",
    "【职业稳定性类】：【填写劣势/无】": "【职业稳定性类】：{{Weakness_Stability}}",
    "【其他待磨合点】：【填写劣势/无】": "【其他待磨合点】：{{Weakness_Other}}",
    "【填写薪资契合度，例：薪资基本契合，提成可协商/差距较大】": "{{SalaryAlignment}}",
    "【填写是否符合招聘要求，例：可快速到岗/到岗时间过晚】": "{{ArrivalDate}}",
    "【填写等级+核心问题/无】": "{{MockPerformance}}",
    "【了解/基本了解/不了解】": "{{ProductKnowledge}}",
    "是否需培训：【是/否】": "是否需培训：{{NeedsTraining}}",
    "【填写细节/无】": "{{OtherDetails}}",
    "评估等级【填写等级】": "评估等级{{FinalGrade}}",
    "与电销岗适配度【高/中/低】": "与电销岗适配度{{SuitabilityLevel}}",
    "核心依据：【填写核心依据】": "核心依据：{{FinalConclusionReason}}",
    "【优先录用/谨慎录用/不予录用】": "{{FinalDecision}}",
    "建议核心依据：【填写依据】": "建议核心依据：{{FinalDecisionReason}}",
}

def fix_template(doc_path):
    """
    修复模板文件中的占位符，将【】替换为{{}}。
    """
    if not os.path.exists(doc_path):
        print(f"Error: Template file not found at {doc_path}")
        return

    # 1. 备份文件
    backup_path = doc_path + '.bak'
    shutil.copy(doc_path, backup_path)
    print(f"Backup of the template created at: {backup_path}")

    doc = Document(doc_path)

    def replace_in_paragraph(paragraph):
        # 为了处理跨越多个run的占位符，我们先构建完整的文本
        full_text = ''.join(run.text for run in paragraph.runs)
        
        modified_text = full_text
        for old, new in replacements.items():
            # 针对表格中的短文本，进行精确匹配或包含匹配
            if old in modified_text:
                # 避免重复替换，例如 "年龄：{{Age}}" 中已经包含了 "年龄"
                if "{{" not in modified_text:
                    modified_text = modified_text.replace(old, new)

        # 如果文本被修改，则用新文本重建段落，同时尽量保留原始格式
        if modified_text != full_text:
            # 清空段落内所有内容
            for run in paragraph.runs:
                run.text = ''
            # 添加带有新文本的run，这会继承段落的样式
            paragraph.add_run(modified_text)

    # 遍历文档的所有部分进行替换
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
    
    # 保存修改后的文档
    doc.save(doc_path)
    print(f"Successfully updated placeholders in: {doc_path}")

if __name__ == "__main__":
    template_file = 'my_template.docx.docx'
    fix_template(template_file)
